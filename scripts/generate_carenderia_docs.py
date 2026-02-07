"""Generate Carenderia Getting Started and User Manual Word documents."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
LOGO_PATH = ROOT / "app" / "static" / "images" / "wvc_logo.png"
VERSION = "1.0"
PREPARED_BY = "SerKai IT Solutions"


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _today_str() -> str:
    return date.today().strftime("%d %b %Y")


def _add_footer(doc: Document) -> None:
    text = f"Version {VERSION} | {_today_str()} | Prepared by: {PREPARED_BY}"
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cover_page(doc: Document, title: str, subtitle: str | None = None) -> None:
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))
        doc.add_paragraph("")
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p_sub = doc.add_paragraph(subtitle)
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    p_meta = doc.add_paragraph(f"Version {VERSION}")
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta2 = doc.add_paragraph(f"Date: {_today_str()}")
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta3 = doc.add_paragraph(f"Prepared by: {PREPARED_BY}")
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def _add_markdownish(doc: Document, md: str) -> None:
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        code_lines = []

    for raw in md.splitlines():
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:]
            doc.add_paragraph(text, style="List Bullet")
            continue
        stripped = line.lstrip()
        if len(stripped) > 3 and stripped[0].isdigit() and stripped[1:3] == ". ":
            doc.add_paragraph(stripped[3:], style="List Number")
            continue
        doc.add_paragraph(line)
    if in_code:
        flush_code()


GETTING_STARTED_CONTENT = """# Carenderia — Getting Started

## Who this guide is for
- Carenderia department users who record daily sales, wages, and expenses.
- Corporate or Admin users who view transactions, trial balance, wages reports, and export PDFs.

## Access requirements
- You must be logged in with a user account.
- Your department must be **Carenderia** or **Corporate** to use Carenderia features.
- Some pages (View Transactions, Monthly Trial Balance, View Wages Report, Export Trial Balance) are available only to **Corporate** or **Admin** role.

## Before you start (setup checklist)
1. Ensure your user account is assigned to **Carenderia** or **Corporate** department.
2. If you are Corporate/Admin: ensure **departments** exist via **Manage Department** (e.g. Carenderia).
3. Add **employees** in **Manage Employee**: name, role, rate per day, and department. These are used for wages.

## Transaction types (what you can record)
- **Daily Sales** — daily collection/sales.
- **Wages** — employee wages for the day.
- **Daily Expense** — general daily expenses.
- **Electric Bill**, **Water Bill**, **Maintenance**, **Mayor's Permit**, **Rental** — fixed or recurring expenses.
- **BIR**, **SSS**, **PAG-IBIG** — government/contributions.
- **Purchases** — purchases with optional reference number and line items (description, qty, unit, unit price, amount).

## Quick workflow (daily recording)
1. Go to **Carenderia → Home** (or use the Carenderia link from the main dashboard).
2. Use **View Transactions** (Corporate/Admin): select a date, add or edit transactions (Daily Sales, Wages, Daily Expense, bills, Purchases, etc.).
3. For **Purchases**, you can enter a reference number and add line items. For **Wages**, wage entries are shown per date with expandable rows.
4. Save your entries. The running balance and daily details update automatically.

## Viewing reports (Corporate/Admin)
- **Monthly Trial Balance** — select a month to see daily collection, total deductions (wages, daily expense, electric, water, maintenance, mayor's permit, rental, BIR, SSS, PAG-IBIG, purchases), and net amount per day.
- **View Wages Report** — view wages by month and by date.
- **Export Trial Balance** — generate a PDF trial balance for a selected month.

## Tips
- Set up employees and their rate per day before recording wages.
- Use **View Transactions** to correct or add entries for any date (Corporate/Admin).
- Daily expense types and amounts can be configured and are included in the trial balance deductions.
"""


USER_MANUAL_CONTENT = """# Carenderia — User Manual

## Overview
The Carenderia module supports daily operations for a canteen/carenderia: recording daily sales, wages, expenses (bills, permits, contributions, purchases), and viewing trial balance and wages reports. Access is controlled by department (Carenderia, Corporate) and role (Admin).

## Access and permissions
- **Carenderia department**: Can access Carenderia Home, Manage Employee, and (where applicable) daily expense settings. Cannot access View Transactions, Monthly Trial Balance, View Wages Report, or Export Trial Balance unless also Admin.
- **Corporate department**: Full access to all Carenderia pages including Manage Department, Manage Employee, View Transactions, Monthly Trial Balance, View Wages Report, and Export Trial Balance.
- **Admin role**: Same as Corporate for Carenderia; can access all Carenderia features regardless of department.

## Home page
- From the main dashboard, click **Carenderia** to open the Carenderia home page.
- The home page shows quick links to:
  - Manage Employee
  - View Transactions (Corporate/Admin)
  - Monthly Trial Balance (Corporate/Admin)
  - View Wages Report (Corporate/Admin)
  - Export Trial Balance (Corporate/Admin)
  - Manage Department (Corporate only)
- Transaction records for recent dates may be summarized on the home page.

## Manage Employee
- **Purpose**: Add, edit, and delete employees used for Carenderia wages.
- **Add employee**: Enter name, role, rate per day, and select department (e.g. Carenderia). Submit.
- **Edit employee**: Use the edit action on a row; update name, role, rate per day, or department; save.
- **Delete employee**: Use the delete action. Deletion may be blocked if the employee is linked to existing wage records.
- Employees listed here are available when recording wages in View Transactions.

## Manage Department (Corporate only)
- **Purpose**: Add, edit, and delete departments (e.g. Carenderia).
- **Add department**: Enter department name; submit. Duplicate names are not allowed.
- **Edit department**: Change the name; save. Ensure no duplicate name with another department.
- **Delete department**: Delete only if the department is not linked to employees or other records.

## View Transactions
- **Purpose**: View and edit all transactions for a selected date. Available to Corporate and Admin.
- **Select date**: Use the date picker to choose the day.
- **Transaction list**: Displays all transactions for that date: type, amount, and for Purchases the reference number and expandable line items (description, qty, unit, unit price, amount). Wages for the date appear with expandable rows showing each wage entry.
- **Add transaction**: Choose transaction type (Daily Sales, Wages, Daily Expense, Electric Bill, Water Bill, Maintenance, Mayor's Permit, Rental, BIR, SSS, PAG-IBIG, Purchases), enter amount. For Purchases you can add reference number and line items; for Wages, entries are managed per date.
- **Edit/Delete**: Use the edit or delete action on a transaction row. Edits apply to the selected transaction; purchase items are tied to that purchase transaction.
- **Running balance**: The page shows a running balance: daily collection minus deductions (wages, daily expense, and all other deduction types). Deductions are listed in ascending order in the deductions table.
- **Daily details**: Daily Details card shows the breakdown of deductions including Daily Expense for the day.

## Transaction types (reference)
- **Daily Sales**: Income; increases daily collection.
- **Wages**: Deduction; sum of wage entries for the date.
- **Daily Expense**: Deduction; amount from daily expense settings or entered for the day.
- **Electric Bill**, **Water Bill**, **Maintenance**, **Mayor's Permit**, **Rental**: Deductions.
- **BIR**, **SSS**, **PAG-IBIG**: Deductions (contributions).
- **Purchases**: Deduction; can have a reference number and multiple line items (description, qty, unit, unit price, amount).

## Monthly Trial Balance
- **Purpose**: View daily collection, total deductions, and net amount for each day in a selected month. Available to Corporate and Admin.
- **Select month**: Choose year and month (e.g. YYYY-MM).
- **Daily details**: For each date you see daily collection, wages, daily expense, electric, water, maintenance, mayor's permit, rental, BIR, SSS, PAG-IBIG, purchases, total deductions, and net amount.
- **Totals**: Month totals for collection and deductions are shown; net = collection minus deductions.

## View Wages Report
- **Purpose**: View wages by month and by date. Available to Corporate and Admin.
- **Select month**: Choose year and month.
- **Report**: Lists each date that has wage entries; expand a date to see employee name, role, rate, and amount. Total wages per date and for the month are shown.

## Export Trial Balance
- **Purpose**: Generate a PDF trial balance for a selected month. Available to Corporate and Admin.
- **Select month**: Choose year and month.
- **Generate PDF**: The system builds a PDF with company header, report title, date range, summary table (Date, Daily Collection, Total Deductions, Net Amount), and detailed transactions by date. Download the file for your records.

## Daily expenses (configuration)
- Default expense types (e.g. Electric Bill, Water Bill, Maintenance, Mayor's Permit, Rental, BIR, SSS, PAG-IBIG) can have default amounts configured.
- These amounts are used in the Monthly Trial Balance and related calculations. Corporate/Carenderia users with access can update or add daily expense types and amounts via the applicable settings or transaction entry.
"""


def build_getting_started_docx() -> Path:
    out_path = DOCS_DIR / "Carenderia_Getting_Started.docx"
    doc = Document()
    _set_base_style(doc)
    _add_cover_page(
        doc,
        title="Carenderia — Getting Started",
        subtitle="Quick guide for daily recording and reports",
    )
    _add_markdownish(doc, GETTING_STARTED_CONTENT)
    _add_footer(doc)
    doc.save(out_path)
    return out_path


def build_user_manual_docx() -> Path:
    out_path = DOCS_DIR / "Carenderia_User_Manual.docx"
    doc = Document()
    _set_base_style(doc)
    _add_cover_page(
        doc,
        title="Carenderia — User Manual",
        subtitle="Full guide to Carenderia features and workflows",
    )
    _add_markdownish(doc, USER_MANUAL_CONTENT)
    _add_footer(doc)
    doc.save(out_path)
    return out_path


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    getting_started = build_getting_started_docx()
    user_manual = build_user_manual_docx()
    print(f"Generated: {getting_started}")
    print(f"Generated: {user_manual}")


if __name__ == "__main__":
    main()
