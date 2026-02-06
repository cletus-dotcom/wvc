from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
LOGO_PATH = ROOT / "app" / "static" / "images" / "wvc_logo.png"
VERSION = "1.0"
PREPARED_BY = "SerKai IT Solutions"


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _today_str() -> str:
    # e.g. 06 Feb 2026
    return date.today().strftime("%d %b %Y")


def _add_footer(doc: Document) -> None:
    text = f"Version {VERSION} | {_today_str()} | Prepared by: {PREPARED_BY}"
    for section in doc.sections:
        footer = section.footer
        # use the first paragraph (Word always has one)
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cover_page(doc: Document, title: str, subtitle: str | None = None) -> None:
    # Logo
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))
        doc.add_paragraph("")

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        p_sub = doc.add_paragraph(subtitle)
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    p_meta = doc.add_paragraph(f"Version {VERSION}")
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta2 = doc.add_paragraph(f"Date: {_today_str()}")
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta3 = doc.add_paragraph(f"Prepared by: {PREPARED_BY}")
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def _add_markdownish(doc: Document, md: str) -> None:
    """
    Minimal Markdown-to-Word renderer (headings, bullets, code blocks).
    This is intentionally simple and predictable for internal manuals.
    """
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        code_lines = []

    for raw in md.splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        # Headings (#, ##, ###)
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        # Bullets
        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:]
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Numbered list (basic: "1. ")
        stripped = line.lstrip()
        if len(stripped) > 3 and stripped[0].isdigit() and stripped[1:3] == ". ":
            doc.add_paragraph(stripped[3:], style="List Number")
            continue

        # Tables in the markdown manual are kept as plain text blocks (simple + reliable)
        doc.add_paragraph(line)

    if in_code:
        flush_code()


def build_user_manual_docx() -> Path:
    md_path = DOCS_DIR / "Construction_User_Manual.md"
    out_path = DOCS_DIR / "Construction_User_Manual.docx"
    md = md_path.read_text(encoding="utf-8")

    doc = Document()
    _set_base_style(doc)
    _add_cover_page(
        doc,
        title="Construction Department — User Manual",
        subtitle="Projects, entries, reports, and corporate processes",
    )
    _add_markdownish(doc, md)
    _add_footer(doc)
    doc.save(out_path)
    return out_path


def build_getting_started_docx() -> Path:
    out_path = DOCS_DIR / "Construction_Getting_Started.docx"

    content = """# Construction — Getting Started

## Who this guide is for
- Construction department users who record daily entries (Materials, Labor, Gasoline, Documents).
- Corporate users who create/manage projects and edit entries.

## Before you start (setup checklist)
1. Make sure your user account is assigned to **Construction** or **Corporate** department.
2. Confirm Construction workers are added in **Manage Workers** and each has a **Rate per day**.
3. Confirm projects exist (Corporate creates projects via **New Project**).

## Daily workflow (most common)
1. Go to **Construction → Home**.
2. Search for your project, then click the project row to open it.
3. Choose **Type of Entry**:
   - Materials
   - Labor (with Overtime)
   - Gasoline
   - Documents
4. Select the **Date** (you’ll see the next invoice number for that date).
5. Fill in the rows, confirm the totals, then click **Submit**.

## Labor overtime (how the system computes it)
- Rate per hour = Rate per day ÷ 8
- Overtime amount = Rate per hour × Overtime hours
- Labor total amount = (Rate per day × Days) + Overtime amount

## Viewing totals and invoices
1. Go to **Construction → View**.
2. Click your project.
3. Use **View By Invoice** and click an invoice row to expand and see the line items.

## Correcting a mistake (Corporate only)
1. Go to **Corporate → Update Project**.
2. Click **Edit Entry** for the project.
3. Expand the invoice, click **Edit** on the line item, update values, and **Save**.

## Tips
- Use the search box on Home/Reports to quickly locate a project.
- Enter the Date first so invoice numbering stays organized by day.
- If you don’t see Corporate/Admin features, your role/department may not have access.

"""

    doc = Document()
    _set_base_style(doc)
    _add_cover_page(
        doc,
        title="Construction — Getting Started",
        subtitle="Quick guide for daily entries and viewing invoices",
    )
    _add_markdownish(doc, content)
    _add_footer(doc)
    doc.save(out_path)
    return out_path


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    manual = build_user_manual_docx()
    getting_started = build_getting_started_docx()
    print(f"Generated: {manual}")
    print(f"Generated: {getting_started}")


if __name__ == "__main__":
    main()

