"""Generate Catering Getting Started and User Manual Word documents."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
LOGO_PATH = ROOT / "app" / "static" / "images" / "wvc_logo.png"
VERSION = "1.0"
PREPARED_BY = "SerKai IT Solutions"


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _today_str() -> str:
    return date.today().strftime("%d %b %Y")


def _add_footer(doc: Document) -> None:
    text = f"Version {VERSION} | {_today_str()} | Prepared by: {PREPARED_BY}"
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cover_page(doc: Document, title: str, subtitle: str | None = None) -> None:
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))
        doc.add_paragraph("")
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p_sub = doc.add_paragraph(subtitle)
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    p_meta = doc.add_paragraph(f"Version {VERSION}")
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta2 = doc.add_paragraph(f"Date: {_today_str()}")
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta3 = doc.add_paragraph(f"Prepared by: {PREPARED_BY}")
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def _add_markdownish(doc: Document, md: str) -> None:
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        code_lines = []

    for raw in md.splitlines():
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:]
            doc.add_paragraph(text, style="List Bullet")
            continue
        stripped = line.lstrip()
        if len(stripped) > 3 and stripped[0].isdigit() and stripped[1:3] == ". ":
            doc.add_paragraph(stripped[3:], style="List Number")
            continue
        doc.add_paragraph(line)
    if in_code:
        flush_code()


GETTING_STARTED_CONTENT = """# Catering — Getting Started

## Who this guide is for
- Catering department staff who manage bookings, add payments, wages, purchases, and miscellaneous expenses per booking.
- Corporate or Admin users who view collectibles, edit transactions, view wages, view balance sheet, and export PDFs.

## Access requirements
- You must be logged in with a user account.
- Your department must be **Catering** or **Corporate** to use Catering features.
- Pastel cards and menus vary by role: **Staff + Catering** can access Manage Bookings, View Collectibles; **Corporate** or **Admin** can access Edit Transaction, View Balance Sheet, View Wages.

## Before you start (setup checklist)
1. Ensure your user account is assigned to **Catering** or **Corporate** department.
2. Ensure **Catering** department exists (via Manage Department).
3. Add **employees** in **Manage Workers**: name, role, rate per day. These are used for wages.
4. Add menu items and equipment in **Manage Menu** and **Manage Equipment** (if applicable).
5. Create or confirm bookings via **Manage Bookings**. Set status to **Confirmed** for bookings to appear on Catering Home.

## Catering Home — Main dashboard
- From the main dashboard, click **Catering** to open Catering Home.
- **Pastel cards** (quick links):
  - **Manage Bookings** — View, edit, and manage catering bookings (Staff+Catering, Admin, Corporate).
  - **View Collectibles** — Bookings with balance or unpaid (Staff+Catering, Admin, Corporate).
  - **Edit Transaction** — List bookings and edit wages, expenses, purchases, miscellaneous (Admin, Corporate).
  - **View Balance Sheet** — Summary and bookings with financial statements (Admin, Corporate).
- **Confirmed Bookings** table: Lists all confirmed bookings. Click a row to expand and choose:
  - **Add Payment** — Record Down Payment, Partial Payment, or Full Payment (with booking amount and running balance).
  - **Add Wages** — Add wage entries per employee (rate/day, days, overtime hours, amount).
  - **Add Purchases** — Add purchase line items (reference PUR-date-###, description, qty, unit, unit price).
  - **Add Miscellaneous** — Add miscellaneous expense.

## Quick workflow (recording per booking)
1. Go to **Catering → Home**.
2. In **Confirmed Bookings**, click a row to expand it.
3. Choose **Add Payment**, **Add Wages**, **Add Purchases**, or **Add Miscellaneous**.
4. Enter the **Date** (required). For purchases, the reference number is auto-generated when you select a date.
5. Fill in the form and save. The booking list and reports update automatically.

## Viewing reports (Corporate/Admin)
- **View Collectibles** — Bookings with balance due; expand rows to see payment history.
- **View Wages** — Wages by month and by date.
- **View Balance Sheet** — Summary card (Total Income, Wages, Other Expenses, Net) plus bookings list (Confirmed/Completed) with expandable financial statements (Income, Expenses, Net per booking). Month filter and search bar available. Export PDF when a month is selected.

## Tips
- Set up employees and their rate per day before recording wages.
- Use **Edit Transaction** to correct or edit expense entries (Corporate/Admin).
- Full Payment with zero running balance updates the booking status to **Completed**.
"""


USER_MANUAL_CONTENT = """# Catering — User Manual

## Overview
The Catering module supports catering operations: managing bookings, recording payments (income), wages, purchases, and miscellaneous expenses per booking, and viewing collectibles, wages reports, and balance sheet. Access is controlled by department (Catering, Corporate) and role (Admin, Staff).

## Access and permissions
- **Catering department (Staff)**: Can access Catering Home, Manage Bookings (no Delete), View Collectibles. Can add Payment, Wages, Purchases, Miscellaneous per booking from Confirmed Bookings.
- **Catering department + Corporate or Admin**: Full access including Edit Transaction, View Wages, View Balance Sheet.
- **Corporate department**: Full access to all Catering pages including Manage Bookings (with Delete), Edit Transaction, View Collectibles, View Wages, View Balance Sheet.
- **Admin role**: Same as Corporate for Catering; can access all Catering features regardless of department.

## Catering Home
- **Entry point**: From the main dashboard, click **Catering**.
- **Pastel cards** (quick links):
  - **Manage Bookings** — List and manage all bookings (Staff+Catering, Admin, Corporate). Staff/Catering cannot delete bookings.
  - **View Collectibles** — Bookings with balance or unpaid; expandable payment history.
  - **Edit Transaction** — List bookings with expandable transactions (Wages, Purchases, Miscellaneous, Expenses); edit button per transaction (Admin, Corporate).
  - **View Balance Sheet** — Summary and bookings with financial statements (Admin, Corporate).
- **Confirmed Bookings** table:
  - Lists bookings with status **Confirmed**.
  - Search bar for live filtering by customer, date, items, contact.
  - Click a row to expand; buttons: Add Payment, Add Wages, Add Purchases, Add Miscellaneous.
  - Each action opens a modal for that booking.

## Manage Bookings
- **Purpose**: Create, edit, and manage catering bookings.
- **Add booking**: Enter requestor name, customer address, contact number, email, event date, event time, items requested. Submit.
- **Edit booking**: Click edit on a row; update fields; save.
- **Delete booking**: Admin/Corporate only. Staff+Catering cannot delete.
- **Status**: Pending, Confirmed, Cancelled, Completed. Set to **Confirmed** for a booking to appear on Catering Home.

## Manage Menu
- **Purpose**: Manage menu items (description, price) used in catering.
- **Add/Edit/Delete**: CRUD operations on menu items.

## Manage Equipment
- **Purpose**: Manage equipment (description, rent price) used in catering.
- **Add/Edit/Delete**: CRUD operations on equipment.

## Manage Workers (Manage Employee)
- **Purpose**: Manage employees (name, role, rate per day) used for Catering wages.
- **Add/Edit/Delete**: CRUD operations on employees in the Catering department.
- Employees listed here appear when adding wages from Confirmed Bookings or Edit Transaction.

## Add Payment (Financial Transaction)
- **Purpose**: Record income (Down Payment, Partial Payment, Full Payment) for a booking.
- **Fields**: Date, Booking (read-only), Booking Amount, Running Balance (read-only), Transaction Description, Transaction Amount, Remarks.
- **Running balance**: Computed as Booking Amount minus existing total paid minus current transaction amount.
- **Full Payment**: When selected, transaction amount auto-fills to the running balance. If running balance becomes zero (or nearly zero), booking status updates to **Completed**.

## Add Wages
- **Purpose**: Add wage entries per employee for a booking (or standalone).
- **Fields**: Date, Employee (select), Rate per Day (auto-filled), Number of Days, Overtime Hours (optional), Overtime Amount (read-only: rate/8 × hours), Amount (read-only: days × rate + overtime).
- **Multiple employees**: Add multiple wage entries; Total Wages Amount is the sum.
- **Saves**: Individual wage entries to catering_wages and a summarized Wages expense to catering_expense linked to the booking (if booking_id provided).

## Add Purchases
- **Purpose**: Add purchase line items for a booking (or standalone).
- **Reference number**: Auto-generated as PUR-YYYY-MM-DD-### when date is selected.
- **Line items**: Description, Qty, Unit, Unit Price, Amount (computed). Add rows as needed.
- **Enter key**: Moves to next column; on Unit Price, Enter adds a new row.
- **Saves**: A Purchases expense with reference_number and CateringPurchaseItem line items linked to the booking.

## Add Miscellaneous
- **Purpose**: Add a miscellaneous expense for a booking.
- **Fields**: Date, Description, Amount, Remarks.

## View Collectibles
- **Purpose**: View bookings with balance due (total_due > total_paid) or unpaid (no transactions yet).
- **Table**: Booking, Event Date, Total Due, Total Paid, Balance. Expandable rows show payment history (date, description, amount).
- Access: Staff+Catering, Admin, Corporate.

## Edit Transaction
- **Purpose**: List all bookings with their transactions; edit expense entries. Admin/Corporate only.
- **Table**: Expandable rows per booking. Each booking shows: Customer, Event Date, Event Time, Status, Transaction count.
- **Expand booking**: Nested table of transactions (Date, Type, Description, Amount, Remarks, Edit button).
- **Purchases and Wages rows**: Expandable to show purchase line items or wage entries.
- **Edit button**: Opens inline form to update Date, Type, Amount, Description, Remarks. Save to update.

## View Wages
- **Purpose**: View wages report by month and by date. Admin/Corporate only.
- **Month filter**: Select month to filter.
- **Report**: Lists each date with wage entries; total per date and for the month.

## View Balance Sheet
- **Purpose**: View summary (Total Income, Total Wages, Other Expenses, Net) and bookings list with financial statements. Admin/Corporate only.
- **Summary card**: Month filter. Four boxes: Total Income, Total Wages, Other Expenses, Net.
- **Bookings card**: Lists Confirmed and Completed bookings. Columns: Booking/Customer, Event Date, Event Time, Status, Total Amount, Income, Expenses, Net. Search bar for live filtering.
- **Expand booking**: Shows Income (payments) table and Expenses table per booking. Booking net at bottom.
- **Export PDF**: Available when a month is selected; generates Balance Sheet PDF.

## Transaction types (reference)
- **Income**: CateringTransaction (Down Payment, Partial Payment, Full Payment) linked to booking_id. booking_amount stores contract total.
- **Wages**: CateringExpense (type Wages) + CateringWage entries (employee, rate, days, overtime, amount).
- **Purchases**: CateringExpense (type Purchases) + CateringPurchaseItem (description, qty, unit, unit_price, amount). Reference PUR-date-###.
- **Miscellaneous**: CateringExpense (type Miscellaneous).
- **Expenses**: CateringExpense (type Expenses) for general expenses.

## Quick Reference
- Add payment per booking: Catering Home → Confirmed Bookings → expand row → Add Payment
- Add wages per booking: Catering Home → Confirmed Bookings → expand row → Add Wages
- Add purchases per booking: Catering Home → Confirmed Bookings → expand row → Add Purchases
- Add miscellaneous per booking: Catering Home → Confirmed Bookings → expand row → Add Miscellaneous
- Manage bookings: Manage Bookings card or Settings → Manage Bookings
- View collectibles: View Collectibles card or Corporate → View Collectibles
- Edit transactions: Edit Transaction card or Corporate → Edit Transaction
- View balance sheet: View Balance Sheet card or Corporate → View Balance Sheet
- View wages: Corporate → View Wages
"""


def build_getting_started_docx() -> Path:
    out_path = DOCS_DIR / "Catering_Getting_Started.docx"
    doc = Document()
    _set_base_style(doc)
    _add_cover_page(
        doc,
        title="Catering — Getting Started",
        subtitle="Quick guide for bookings, payments, wages, and reports",
    )
    _add_markdownish(doc, GETTING_STARTED_CONTENT)
    _add_footer(doc)
    doc.save(out_path)
    return out_path


def build_user_manual_docx() -> Path:
    out_path = DOCS_DIR / "Catering_User_Manual.docx"
    doc = Document()
    _set_base_style(doc)
    _add_cover_page(
        doc,
        title="Catering — User Manual",
        subtitle="Full guide to Catering features and workflows",
    )
    _add_markdownish(doc, USER_MANUAL_CONTENT)
    _add_footer(doc)
    doc.save(out_path)
    return out_path


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    getting_started = build_getting_started_docx()
    user_manual = build_user_manual_docx()
    print(f"Generated: {getting_started}")
    print(f"Generated: {user_manual}")


if __name__ == "__main__":
    main()
